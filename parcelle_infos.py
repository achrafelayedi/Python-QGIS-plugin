# -*- coding: utf-8 -*-
"""
/***************************************************************************
 ParcelleInfos
                                 A QGIS plugin
 Déterminer les parcelles susceptibles d'accueillir des projets de dérogation
                              -------------------
        begin                : 2018-07-08
        git sha              : $Format:%H$
        copyright            : (C) 2018 by EL AYEDI Achraf 
        email                : achrafelayediz@gmail.com 
 ***************************************************************************/

/***************************************************************************
 *                                                                         *
 *   This program is free software; you can redistribute it and/or modify  *
 *   it under the terms of the GNU General Public License as published by  *
 *   the Free Software Foundation; either version 2 of the License, or     *
 *   (at your option) any later version.                                   *
 *                                                                         *
 ***************************************************************************/
"""

# import pour l'encodage utf-8
import sys

reload(sys)
sys.setdefaultencoding("utf-8")

import os
import qgis
import PyQt4
from qgis.utils import *
from qgis.gui import *
from PyQt4.QtGui import QImage

from PyQt4.QtCore import QSettings, QTranslator, qVersion, QCoreApplication, QSize
from PyQt4.QtGui import *
from qgis.core import *
from os.path import expanduser

# Initialize Qt resources from file resources.py
from qgis._core import QgsMessageLog

import resources
# Import the code for the dialog
from parcelle_infos_dialog import ParcelleInfosDialog
import os.path

# les variables globales

bufferLayer = None
msgerr = []
lcpt = -1
lcolor = []

class ParcelleInfos:
    """QGIS Plugin Implementation."""
    def __init__(self, iface):
        """Constructor.

        :param iface: An interface instance that will be passed to this class
            which provides the hook by which you can manipulate the QGIS
            application at run time.
        :type iface: QgisInterface
        """
        # Save reference to the QGIS interface
        self.iface = iface
        # initialize plugin directory
        self.plugin_dir = os.path.dirname(__file__)
        # initialize locale
        locale = QSettings().value('locale/userLocale')[0:2]
        locale_path = os.path.join(
            self.plugin_dir,
            'i18n',
            'ParcelleInfos_{}.qm'.format(locale))

        if os.path.exists(locale_path):
            self.translator = QTranslator()
            self.translator.load(locale_path)

            if qVersion() > '4.3.3':
                QCoreApplication.installTranslator(self.translator)

        # Create the dialog (after translation) and keep reference
        self.dlg = ParcelleInfosDialog()

        # Declare instance attributes
        self.actions = []
        self.menu = self.tr(u'&Parcelle Infos')
        # We are going to let the user set this up in a future iteration
        self.toolbar = self.iface.addToolBar(u'ParcelleInfos')
        self.toolbar.setObjectName(u'ParcelleInfos')
       

    # les fonctions créées

    # la fonction qui affiche un message de validation lors de fermeture de programme
    def closeEvent(self,event):
        box = QMessageBox()
        box.setIcon(QMessageBox.Question)
        box.setWindowTitle('MESSAGE')
        box.setText("\nÊtes-vous sûr de vouloir quitter le programme ?".decode('utf-8'))
        box.setStandardButtons(QMessageBox.Yes|QMessageBox.No)
        buttonY = box.button(QMessageBox.Yes)
        buttonY.setText('Oui')
        buttonN = box.button(QMessageBox.No)
        buttonN.setText('Non')

        ph = self.dlg.geometry().height()
        px = self.dlg.geometry().x()
        py = self.dlg.geometry().y()
        dh = self.dlg.height()
        box.move(px+60, py+ph-dh+250)
        box.exec_()

        if box.clickedButton() == buttonY:
            self.reinitialiser()
            event.accept()
        else:
            event.ignore()

    # la fonction qui sert à reinitialiser les champs du programme
    def reinitialiser(self):
        global bufferLayer
        global lcolor

        lcolor = []
        self.dlg.prjsdero.hide()
        self.dlg.frame_3.hide()
        self.dlg.tableWidget.hide()

        self.dlg.assfonc.hide()
        self.dlg.frame_4.hide()
        self.dlg.tableWidget_2.hide()

        self.dlg.exporter.setEnabled(False)

        self.dlg.ldpe.hide()
        self.dlg.vdpe.hide()
        self.dlg.ldf.hide()
        self.dlg.vdf.hide()
        self.dlg.lc.hide()
        self.dlg.vc.hide()
        self.dlg.ldp.hide()
        self.dlg.vdp.hide()
        self.dlg.ldc.hide()
        self.dlg.vdc.hide()

        QLineEditBlanc = "QLineEdit { background: rgb(255, 255, 255); }"
        self.dlg.x.clear()
        self.dlg.x.setStyleSheet(QLineEditBlanc)

        self.dlg.y.clear()
        self.dlg.y.setStyleSheet(QLineEditBlanc)

        self.dlg.buffer.clear()
        self.dlg.buffer.setStyleSheet(QLineEditBlanc)

        try:
            QgsMapLayerRegistry.instance().removeMapLayer(bufferLayer.id())
            bufferLayer = None
        except:
            pass

    # la fonction qui permet de prendre les informations sur les projets de dérogations intersectés avec le buffer
    def lireCouchePrjDerog(self):
        global bufferLayer
        lignes = []
        coucheDerogation = QgsMapLayerRegistry.instance().mapLayersByName("Derogation_central_13_avril")[0]

        # calcule des colonnes
        colonnes = [field.name() for field in coucheDerogation.pendingFields() ]
        # obtenir les lignes des projets dérogés intersectés avec le buffer
        for fbuffer in bufferLayer.getFeatures():
            cands = coucheDerogation.getFeatures(QgsFeatureRequest().setFilterRect(fbuffer.geometry().boundingBox()))
    
            for area_feature in cands:
                if fbuffer.geometry().intersects(area_feature.geometry()):
                    if area_feature["Region"]:
                        lignes.append(area_feature.attributes())

        return lignes, colonnes
                
    # la fonction qui sert à remplir le tableau des projets de dérogations
    def remplirTabPrjDerog(self):
        global msgerr

        # obtenir la liste des lignes et des colonnes des projets dérogés
        l,c = self.lireCouchePrjDerog()
        
        if len(l):
            # le tableau
            self.dlg.prjsdero.show()
            self.dlg.frame_3.show()
            self.dlg.tableWidget.show()
            self.dlg.tableWidget.setRowCount(len(l))
            self.dlg.tableWidget.setColumnCount(len(c))

            # definir les labels
            self.dlg.tableWidget.setHorizontalHeaderLabels(c)

            # remplir les données
            for il in range(0, len(l)):
                for ic in range(0, len(c)):
                    if "QDate" in str(l[il][ic]):
                            l[il][ic] = l[il][ic].toString("yyyy-MM-dd")
                    cellule = QTableWidgetItem(str(l[il][ic]))
                    self.dlg.tableWidget.setItem(il, ic, cellule)
                    cellule.setBackgroundColor(QColor(255, 234, 0))

        else:
            msgerr.append(1)
            self.dlg.tableWidget.setRowCount(0)
            
    # la fonction qui permet de prendre les informations sur les assiettes foncières intersectés avec le buffer
    def lireCouchesAssFonc(self):
        global bufferLayer

        coucheDomainePriveEtat = QgsMapLayerRegistry.instance().mapLayersByName("DOMIANE_PRIVE_ETAT")[0]

        # calcule des colonnes
        colonnes = [field.name() for field in coucheDomainePriveEtat.pendingFields() ]

        # récupérer les lignes de chaque assiette foncière intersectée avec le buffer
        def lireToutesCouchesAssFonc(nomCouche):
            global bufferLayer
            lignes = []
            coucheDomainePriveEtat = QgsMapLayerRegistry.instance().mapLayersByName(nomCouche)[0]
            # obtenir les lignes des assiettes foncières intersectés avec le buffer
            for fbuffer in bufferLayer.getFeatures():
                cands = coucheDomainePriveEtat.getFeatures(QgsFeatureRequest().setFilterRect(fbuffer.geometry().boundingBox()))
    
                for area_feature in cands:
                    if fbuffer.geometry().intersects(area_feature.geometry()):
                        lignes.append(area_feature.attributes())

            return lignes


        return lireToutesCouchesAssFonc("DOMIANE_PRIVE_ETAT"),lireToutesCouchesAssFonc("DOMAINE_PUBLIC"),lireToutesCouchesAssFonc("DOMAINE_FORESTIER"),lireToutesCouchesAssFonc("DOMAINE_COMMUNAL"),lireToutesCouchesAssFonc("COLLECTIF"), colonnes
    

    # la fonction qui sert à remplir le tableau des assiettes foncières
    def remplirTabAssFonc(self):
        global msgerr
        global lcpt

        # obtenir la liste des lignes et des colonnes des assiettes foncières
        dpe,dp,df,dc,co,c = self.lireCouchesAssFonc()

        l = len(dpe)+len(dp)+len(df)+len(dc)+len(co)
        if l:
            # le tableau
            self.dlg.assfonc.show()
            self.dlg.frame_4.show()
            self.dlg.tableWidget_2.show()
            self.dlg.tableWidget_2.setRowCount(l)
            self.dlg.tableWidget_2.setColumnCount(len(c))

            # definir les labels
            self.dlg.tableWidget_2.setHorizontalHeaderLabels(c)

            # rgb vers hex
            def rgb2hex(r, g, b):
                return "{:02x}{:02x}{:02x}".format(r,g,b)
            
            # fonction locale pour remplir les lignes du tablau
            def remplirTab(l, c, r, g, b):
                global lcpt
                global lcolor
                for il in range(0, len(l)):
                    lcpt += 1
                    for ic in range(0, len(c)):
                        if "QDate" in str(l[il][ic]):
                            l[il][ic] = l[il][ic].toString("yyyy-MM-dd")
                        cellule = QTableWidgetItem(str(l[il][ic]))
                        self.dlg.tableWidget_2.setItem(lcpt, ic, cellule)
                        cellule.setBackgroundColor(QColor(r, g, b))
                    tmpcolor = rgb2hex(r, g, b)
                    lcolor.append(tmpcolor)

            remplirTab(dpe, c, 206, 147, 216)
            remplirTab(dp, c, 77, 208, 225)
            remplirTab(df, c, 129, 199, 132)
            remplirTab(dc, c, 255, 183, 77)
            remplirTab(co, c, 161, 136, 127)
            lcpt = -1

            # l'affichage du pourcentage de chaque assiettes

            def calc(val, l):
                return str(round(float(val*100)/float(l), 2))+" %"

            self.dlg.ldpe.show()
            self.dlg.vdpe.setText(calc(len(dpe), l))
            self.dlg.vdpe.show()

            self.dlg.ldf.show()
            self.dlg.vdf.setText(calc(len(df), l))
            self.dlg.vdf.show()

            self.dlg.lc.show()
            self.dlg.vc.setText(calc(len(co), l))
            self.dlg.vc.show()

            self.dlg.ldp.show()
            self.dlg.vdp.setText(calc(len(dp), l))
            self.dlg.vdp.show()

            self.dlg.ldc.show()
            self.dlg.vdc.setText(calc(len(dc), l))
            self.dlg.vdc.show()

        else:
            msgerr.append(2)
            self.dlg.tableWidget_2.setRowCount(0)

    def messageErreur(self):
        global msgerr
        global lcolor

        if msgerr:
            self.dlg.exporter.setEnabled(False)
            lcolor = []
            box = QMessageBox()
            box.setIcon(QMessageBox.Warning)
            box.setWindowTitle('MESSAGE')

            if len(msgerr) == 2:
                box.setText("\nCette zone ne contient pas ni de projets dérogés\nni des assiettes foncières qui dépassent 1 Ha !".decode('utf-8'))
            elif 1 in msgerr:
                box.setText("\nCette zone ne contient pas de projets dérogés !".decode('utf-8'))
            else:
                box.setText("\nCette zone ne contient pas des assiettes foncières\nqui dépassent 1 Ha !".decode('utf-8'))
            
            ph = self.dlg.geometry().height()
            px = self.dlg.geometry().x()
            py = self.dlg.geometry().y()
            dh = self.dlg.height()
            box.move(px+60, py+ph-dh+250)
            box.exec_()

            msgerr = []
        else:
            self.dlg.exporter.setEnabled(True)

    # la fonction qui sert à executer le programme
    def executer(self):
        global bufferLayer
        testvar = None

        # définir le couleur rouge des champs vides
        QLineEditRouge = "QLineEdit { background: rgb(247, 54, 76); }"
        # la validation des valeurs des champs
        if not self.dlg.x.text():
            self.dlg.x.setStyleSheet(QLineEditRouge)
            testvar = True
        else:
            x = float(self.dlg.x.text())

        if not self.dlg.y.text():
            self.dlg.y.setStyleSheet(QLineEditRouge)
            testvar = True
        else:
            y = float(self.dlg.y.text())

        if not self.dlg.buffer.text():
            self.dlg.buffer.setStyleSheet(QLineEditRouge)
            testvar = True
        else:
            buffer = int(self.dlg.buffer.text())

        # si tous les champs sont remplis
        if not testvar:

            if bufferLayer:
                # reinitialiser le buffer
                QgsMapLayerRegistry.instance().removeMapLayer(bufferLayer.id())
                self.dlg.prjsdero.hide()
                self.dlg.frame_3.hide()
                self.dlg.tableWidget.hide()

                self.dlg.assfonc.hide()
                self.dlg.frame_4.hide()
                self.dlg.tableWidget_2.hide()

                self.dlg.ldpe.hide()
                self.dlg.vdpe.hide()
                self.dlg.ldf.hide()
                self.dlg.vdf.hide()
                self.dlg.lc.hide()
                self.dlg.vc.hide()
                self.dlg.ldp.hide()
                self.dlg.vdp.hide()
                self.dlg.ldc.hide()
                self.dlg.vdc.hide()
            
            # la création de la couche vecteur du buffer qui contient comme centre le point choisi

            # obtenir le CRS (Coordinate Reference System) du projet 
            canvas = self.iface.mapCanvas()
            mapRenderer = canvas.mapRenderer()
            crs=mapRenderer.destinationCrs()

            bufferLayer = QgsVectorLayer("Polygon?crs="+str(crs.toWkt()), "LA ZONE D'ETUDE", "memory")
            feature = QgsFeature()
            feature.setGeometry( QgsGeometry.fromPoint(QgsPoint(x, y)).buffer(buffer, 100))
            provider = bufferLayer.dataProvider()
            bufferLayer.startEditing()
            provider.addFeatures([feature])
            
            symbols = bufferLayer.rendererV2().symbols()
            symbol = symbols[0]
            symbol.setColor(QColor.fromRgb(0, 255, 230))
            
            bufferLayer.setLayerTransparency(60)
            
            bufferLayer.commitChanges()
            QgsMapLayerRegistry.instance().addMapLayers([bufferLayer])

            # toujours mettre la couche buffer en haut
            bridge = self.iface.layerTreeCanvasBridge() 
            order = bridge.customLayerOrder()
            order.insert( 0, order.pop( order.index( bufferLayer.id() ) ) )
            bridge.setCustomLayerOrder( order )

            extent = bufferLayer.extent()
            canvas.setExtent(extent)

            
            # remplir le tableau des projets de dérogation
            self.remplirTabPrjDerog()
            # remplir le tableau des assiettes foncières
            self.remplirTabAssFonc()

            # vérifier s'il n'existe pas de projets de dérogation et/ou des assiettes foncières 
            self.messageErreur()
              
    # changer les couleurs des couches de projet
    def changerCouleurCouches(self, NomCouche, r, g, b):
        couche = QgsMapLayerRegistry.instance().mapLayersByName(NomCouche)[0]
        symbols = couche.rendererV2().symbols()
        symbol = symbols[0]
        symbol.setColor(QColor.fromRgb(r, g, b))
        self.iface.mapCanvas().refreshAllLayers() 
        self.iface.legendInterface().refreshLayerSymbology(couche)
    
    def exporter(self):
        global lcolor
        import docx
        from docx.shared import Inches, Pt, RGBColor

        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml

        from docx.enum.table import WD_TABLE_ALIGNMENT

        from docx import Document

        # PARAMETERS
        mainPath = str(expanduser("~"))+"/"
        filename = 'tmp'
        imageType = "png"
        dpi = 300

        map_settings = self.iface.mapCanvas().mapSettings()
        c = QgsComposition(map_settings)
        #c.setPaperSize(400, 160)
        c.setPrintResolution(dpi)

        #set page background to transparent
        transparent_fill =QgsFillSymbolV2.createSimple({ 'outline_style': 'no', 'style': 'no'})
        c.setPageStyleSymbol( transparent_fill )

        x, y = 0, 0
        w, h = c.paperWidth(), c.paperHeight()
        composerMap = QgsComposerMap(c, x ,y, w, h)
        composerMap.setBackgroundEnabled(False)
        c.addItem(composerMap)

        mapRenderer = self.iface.mapCanvas().mapRenderer()
        composerLabel = QgsComposerLabel(c)

        composerLabel.adjustSizeToText()
        c.addItem(composerLabel)

        legend = QgsComposerLegend(c)
        legend.setTitle("             LÉGENDE".decode('utf-8'))
        legend.model().setLayerSet(mapRenderer.layerSet())
        c.addItem(legend)

        dpmm = dpi / 25.4
        width = int(dpmm * c.paperWidth())
        height = int(dpmm * c.paperHeight())

        # create output image and initialize it
        image = QImage(QSize(width, height), QImage.Format_ARGB32)
        image.setDotsPerMeterX(dpmm * 1000)
        image.setDotsPerMeterY(dpmm * 1000)
        #image.fill(Qt.transparent)

        imagePainter = QPainter(image)

        c.setPlotStyle(QgsComposition.Print)
        c.renderPage( imagePainter, 0 )
        imagePainter.end()

        imageFilename =  mainPath + filename + '.' + imageType
        image.save(imageFilename, imageType)

        #----------------------------------------------

        # créer un document
        document = docx.Document()
        # définir les margines du document
        section = document.sections[0]
        section.right_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        document.add_picture(str(mainPath)+str(filename)+"."+str(imageType), width=Inches(7.25))
        
        # nombres de lignes et des colonnes du tableau des projets dérogés 
        l = self.dlg.tableWidget.rowCount()
        c = self.dlg.tableWidget.columnCount()

        cprjdrg = ["Region", "Province", "Commune", "Agence", "Superficie", "Forme_MOA", "Nature_pro"]
        
        # lire les valeurs du tableau des projets dérogés 
        def lireTabDerog(l, c, cprjdrg):
            data = []
            donnees = []
    
            for il in range(l):
                for ic in range(c):
                    if self.dlg.tableWidget.horizontalHeaderItem(ic).text() in cprjdrg:
                        data.append(self.dlg.tableWidget.item( il, ic).text())
                donnees.append(data)
                data = []
            return donnees

        document.add_heading("                              LES PROJETS DE DÉROGATIONS\n".decode('utf-8'))
        # ajouter le tableau des projets dérogés 
        table = document.add_table(rows=1, cols=7)
        
        # ajouter les noms des colonnes
        heading_cells = table.rows[0].cells
        for i in range(len(cprjdrg)):
            Nombre_text_formatted = heading_cells[i].paragraphs[0].add_run(cprjdrg[i])
            Nombre_text_formatted.font.size = Pt(8)
            Nombre_text_formatted.font.name = 'Verdana'
            

        # remplir le tableau des projets dérogés
        prjsdrj = lireTabDerog(l, c, cprjdrg)
        for d in range(len(prjsdrj)):
            cells = table.add_row().cells
            for i in range(7):
                #cells[i].text = str(prjsdrj[d][i])
                row = table.rows[d+1]
                Nombre_text_formatted = row.cells[i].paragraphs[0].add_run(prjsdrj[d][i])
                Nombre_text_formatted.font.size = Pt(8)
                Nombre_text_formatted.bold = False

                shading_elm = parse_xml(r'<w:shd {} w:fill="ffea00"/>'.format(nsdecls('w')))
                cells[i]._tc.get_or_add_tcPr().append(shading_elm)
        
        table.style = 'Colorful List'

        #----------------------------------------------------------------

        # nombres de lignes et des colonnes du tableau des assiettes foncières
        l = self.dlg.tableWidget_2.rowCount()
        c = self.dlg.tableWidget_2.columnCount()

        cassfon = ["REGIME_FON", "NATURE_FON", "CERCLE", "COMMUNE", "STATUT_FON", "SUPERFICIE", "NATURE_OCC"]

        # lire les valeurs du tableau des assiettes foncières 
        def lireTabAssFonc(l, c, cassfon):
            data = []
            donnees = []
        
            for il in range(l):
                for ic in range(c):
                    if self.dlg.tableWidget_2.horizontalHeaderItem(ic).text() in cassfon:
                        data.append(self.dlg.tableWidget_2.item( il, ic).text())
                donnees.append(data)
                data = []
            return donnees

        document.add_heading("                                  LES ASSIETTES FONCIÈRES\n".decode('utf-8'))
        # ajouter le tableau des assiettes foncières
        table2 = document.add_table(rows=1, cols=7)
        # ajouter les noms des colonnes
        heading_cells = table2.rows[0].cells
        for i in range(len(cassfon)):
            #heading_cells[i].text = str(cassfon[i])
            Nombre_text_formatted = heading_cells[i].paragraphs[0].add_run(cassfon[i])
            Nombre_text_formatted.font.size = Pt(8)
            Nombre_text_formatted.font.name = 'Verdana'

        # remplir le tableau des assiettes foncières 
        assfon = lireTabAssFonc(l, c, cassfon)

        for d in range(len(assfon)):
            cells = table2.add_row().cells
            for i in range(7):
                row = table2.rows[d+1]
                Nombre_text_formatted = row.cells[i].paragraphs[0].add_run(assfon[d][i])
                Nombre_text_formatted.font.size = Pt(8)
                Nombre_text_formatted.bold = False
                if lcolor[d] == "81c784":
                    shading_elm = parse_xml(r'<w:shd {} w:fill="81c784"/>'.format(nsdecls('w')))
                elif lcolor[d] == "4dd0e1":
                    shading_elm = parse_xml(r'<w:shd {} w:fill="4dd0e1"/>'.format(nsdecls('w')))
                elif lcolor[d] == "ce93d8":
                    shading_elm = parse_xml(r'<w:shd {} w:fill="ce93d8"/>'.format(nsdecls('w')))
                elif lcolor[d] == "ffb74d":
                    shading_elm = parse_xml(r'<w:shd {} w:fill="ffb74d"/>'.format(nsdecls('w')))
                elif lcolor[d] == "a1887f":
                    shading_elm = parse_xml(r'<w:shd {} w:fill="a1887f"/>'.format(nsdecls('w')))

                cells[i]._tc.get_or_add_tcPr().append(shading_elm)
        
        table2.style = 'Colorful List'

        #----------------------------------------------------------------

        run = document.add_paragraph().add_run("\n\n                            DOMIANE_PRIVE_ETAT : "+str(self.dlg.vdpe.text()) )
        font = run.font
        font.color.rgb = RGBColor(206, 147, 216)
        run.bold = True

        run = document.add_paragraph().add_run("                                                                DOMAINE_FORESTIER : "+str(self.dlg.vdf.text()) )
        font = run.font
        font.color.rgb = RGBColor(129, 199, 132)
        run.bold = True

        run = document.add_paragraph().add_run("                            DOMAINE_COMMUNAL : "+str(self.dlg.vdc.text()) )
        font = run.font
        font.color.rgb = RGBColor(255, 183, 77)
        run.bold = True
        
        run = document.add_paragraph().add_run("                                                                DOMAINE_PUBLIC : "+str(self.dlg.vdp.text()) )
        font = run.font
        font.color.rgb = RGBColor(77, 208, 225)
        run.bold = True

        run = document.add_paragraph().add_run("                            COLLECTIF : "+str(self.dlg.vc.text()) )
        font = run.font
        font.color.rgb = RGBColor(161, 136, 127)
        run.bold = True

        document.save(str(expanduser("~"))+"/rapport.docx")
        os.remove(str(mainPath)+str(filename)+"."+str(imageType))

        lcolor = []

        box = QMessageBox()
        box.setIcon(QMessageBox.Information)
        box.setWindowTitle('MESSAGE')
        box.setText("\nLe rapport a été créé avec succès dans :\n"+str(expanduser("~"))+"/rapport.docx".decode('utf-8'))
        ph = self.dlg.geometry().height()
        px = self.dlg.geometry().x()
        py = self.dlg.geometry().y()
        dh = self.dlg.height()
        box.move(px+60, py+ph-dh+250)
        box.exec_()

    def tr(self, message):
        # noinspection PyTypeChecker,PyArgumentList,PyCallByClass
        return QCoreApplication.translate('ParcelleInfos', message)

    def add_action(
            self,
            icon_path,
            text,
            callback,
            enabled_flag=True,
            add_to_menu=True,
            add_to_toolbar=True,
            status_tip=None,
            whats_this=None,
            parent=None):


        # Create the dialog (after translation) and keep reference
        self.dlg = ParcelleInfosDialog()

        icon = QIcon(icon_path)
        action = QAction(icon, text, parent)
        action.triggered.connect(callback)
        action.setEnabled(enabled_flag)

        if status_tip is not None:
            action.setStatusTip(status_tip)

        if whats_this is not None:
            action.setWhatsThis(whats_this)

        if add_to_toolbar:
            self.toolbar.addAction(action)

        if add_to_menu:
            self.iface.addPluginToVectorMenu(
                self.menu,
                action)

        self.actions.append(action)

        return action

    def initGui(self):
        """Create the menu entries and toolbar icons inside the QGIS GUI."""

        icon_path = ':/plugins/ParcelleInfos/icon.png'
        self.add_action(
            icon_path,
            text=self.tr(u''),
            callback=self.run,
            parent=self.iface.mainWindow())           
        # les boutons utilisés
        # le bouton réinitialiser
        self.dlg.reinitialiser.pressed.connect(self.reinitialiser)
        # le bouton exécuter
        self.dlg.executer.pressed.connect(self.executer)
        # le bouton fermer
        self.dlg.closeEvent = self.closeEvent
        # le bouton exporter
        self.dlg.exporter.pressed.connect(self.exporter)
        # cliquer sur une ligne du tableau des projets dérogés 
        self.dlg.tableWidget.clicked.connect(self.cliquerLignePrjDero)
        # cliquer sur une ligne du tableau des projets dérogés 
        self.dlg.tableWidget_2.clicked.connect(self.cliquerLigneAssFonc)

    def cliquerLigneAssFonc(self):
        tmp = []
        layer1 = QgsMapLayerRegistry.instance().mapLayersByName("DOMIANE_PRIVE_ETAT")[0]
        layer2 = QgsMapLayerRegistry.instance().mapLayersByName("DOMAINE_PUBLIC")[0]
        layer3 = QgsMapLayerRegistry.instance().mapLayersByName("DOMAINE_FORESTIER")[0]
        layer4 = QgsMapLayerRegistry.instance().mapLayersByName("DOMAINE_COMMUNAL")[0]
        layer5 = QgsMapLayerRegistry.instance().mapLayersByName("COLLECTIF")[0]
        
        atable = self.dlg.tableWidget_2

        for i in atable.selectionModel().selectedRows():
            ID = atable.item(i.row(),0).text()
            SUP = atable.item(i.row(),10).text()
        
        selection = layer1.getFeatures(QgsFeatureRequest(QgsExpression("\"OBJECTID\" = '"+str(ID)+"' AND \"SUPERFICIE\" = '"+str(SUP)+"'")))
        for feature in selection:
            tmp.append(feature.id())
            layer1.setSelectedFeatures(tmp)
            box = layer1.boundingBoxOfSelected()

        selection = layer2.getFeatures(QgsFeatureRequest(QgsExpression("\"OBJECTID\" = '"+str(ID)+"' AND \"SUPERFICIE\" = '"+str(SUP)+"'")))
        for feature in selection:
            tmp.append(feature.id())
            layer2.setSelectedFeatures(tmp)
            box = layer2.boundingBoxOfSelected()

        selection = layer3.getFeatures(QgsFeatureRequest(QgsExpression("\"OBJECTID\" = '"+str(ID)+"' AND \"SUPERFICIE\" = '"+str(SUP)+"'")))
        for feature in selection:
            tmp.append(feature.id())
            layer3.setSelectedFeatures(tmp)
            box = layer3.boundingBoxOfSelected()

        selection = layer4.getFeatures(QgsFeatureRequest(QgsExpression("\"OBJECTID\" = '"+str(ID)+"' AND \"SUPERFICIE\" = '"+str(SUP)+"'")))
        for feature in selection:
            tmp.append(feature.id())
            layer4.setSelectedFeatures(tmp)
            box = layer4.boundingBoxOfSelected()
        
        selection = layer5.getFeatures(QgsFeatureRequest(QgsExpression("\"OBJECTID\" = '"+str(ID)+"' AND \"SUPERFICIE\" = '"+str(SUP)+"'")))
        for feature in selection:
            tmp.append(feature.id())
            layer5.setSelectedFeatures(tmp)
            box = layer5.boundingBoxOfSelected()

        
        self.iface.mapCanvas().setExtent(box)
        self.iface.mapCanvas().refresh()
        tmp = []


    def cliquerLignePrjDero(self):
        tmp = []
        layer = QgsMapLayerRegistry.instance().mapLayersByName("Derogation_central_13_avril")[0]
        atable = self.dlg.tableWidget

        for i in atable.selectionModel().selectedRows():
            ID = atable.item(i.row(),8).text()
        
        selection = layer.getFeatures(QgsFeatureRequest(QgsExpression("\"SUPERFICIE\" = '"+str(ID)+"'")))
        for feature in selection:
            tmp.append(feature.id())

        layer.setSelectedFeatures(tmp)
        box = layer.boundingBoxOfSelected()
        self.iface.mapCanvas().setExtent(box)
        self.iface.mapCanvas().refresh()
        tmp = []
        
    def unload(self):
        """Removes the plugin menu item and icon from QGIS GUI."""
        for action in self.actions:
            self.iface.removePluginVectorMenu(
                self.tr(u'&Parcelle Infos'),
                action)
            self.iface.removeToolBarIcon(action)
        # remove the toolbar
        del self.toolbar

    def run(self):
        """Run method that performs all the real work"""
        
        self.changerCouleurCouches("DOMIANE_PRIVE_ETAT", 206, 147, 216)
        self.changerCouleurCouches("DOMAINE_PUBLIC", 77, 208, 225)
        self.changerCouleurCouches("DOMAINE_FORESTIER", 129, 199, 132)
        self.changerCouleurCouches("DOMAINE_COMMUNAL", 255, 183, 77)
        self.changerCouleurCouches("COLLECTIF", 161, 136, 127)
        self.changerCouleurCouches("Derogation_central_13_avril", 255, 234, 0)
        # reinitialiser le programmme
        self.reinitialiser()
        # la validation des champs
        # validateur de la valeur de x
        self.dlg.x.setValidator(QDoubleValidator(-0.99,-99.99,2))
        # validateur de la valeur de y
        self.dlg.y.setValidator(QDoubleValidator(-0.99,-99.99,2))
        # validateur de la valeur du buffer
        self.dlg.buffer.setValidator(QIntValidator())
        self.dlg.buffer.setMaxLength(9)

        # show the dialog
        # fixer la taille du Dialog
        self.dlg.setFixedSize(530, 695)
        self.dlg.show()
        # Run the dialog event loop
        result = self.dlg.exec_()
        # See if OK was pressed
        if result:
            # Do something useful here - delete the line containing pass and
            # substitute with your code.
            pass
